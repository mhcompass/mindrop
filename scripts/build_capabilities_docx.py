"""Generate a branded YourCompass 'AI Ops Platform — Capabilities' Word doc
from the capability catalogue in src/model/plan3.ts.

Output: ../YourCompass_AIOps_Capabilities.docx  (one level up from aiops).
Run:    python3 scripts/build_capabilities_docx.py
"""
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.normpath(os.path.join(HERE, "..", "..", "YourCompass_AIOps_Capabilities.docx"))
LOGO = os.path.normpath(
    os.path.join(HERE, "..", "..", "yc-itsm", "frontend", "public", "logo-blue.png")
)

# Brand palette
BLUE = RGBColor(0x25, 0x63, 0xEB)
NAVY = RGBColor(0x1F, 0x29, 0x37)
GREY = RGBColor(0x6B, 0x72, 0x80)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
BODY = RGBColor(0x37, 0x41, 0x51)

# ── Capability catalogue (mirrors src/model/plan3.ts) ───────────────────
DOMAINS = [
    "Service management (ITIL)",
    "Assistant and knowledge",
    "Platform and operations",
    "System integrations",
]

AVAILABLE = {
    "Service management (ITIL)": [
        "Incident management — create, triage, assign, work-note and transition, with SLA timers and knowledge citations",
        "Problem management — records, root cause and lifecycle, linked to their incidents",
        "Change management — risk classification with a CAB approval gate on high-risk changes",
        "CAB calendar — scheduled changes with conflict flags",
        "Service catalog and citizen portal — submit a request and track it",
        "Asset inventory (CMDB) — assets and sites with filtering",
        "Cross-practice linking — connect incidents, problems and changes",
    ],
    "Assistant and knowledge": [
        "Compass assistant — chat with tool use across all of the above",
        "Knowledge base with citations — folder-scoped question answering with sources",
        "MyLocker document workspace — templates, with English/Arabic document analysis",
        "Reports — natural-language and scheduled reports",
        "Voice — speech-to-text and text-to-speech in English and Arabic",
        "Execution trace — step-by-step view of each assistant turn",
    ],
    "Platform and operations": [
        "Multi-tenant profiles and branding — switch between Default, Dubai Police and Ministry of Education, with data kept separate per tenant",
        "Governance — assistant kill-switch and full audit log",
        "Approvals — email approve/reject with secure tokens",
        "Directory operations — user lookup, group audit and password reset",
        "Device and compliance views (SCCM)",
        "Operator web console and REST API",
    ],
    "System integrations": [
        "Microsoft Entra ID — single sign-on (OIDC)",
        "Microsoft Graph — devices, mail and approvals",
        "Active Directory / LDAP — directory operations",
        "Microsoft Exchange — mailboxes and outbound email",
        "SharePoint — knowledge-base document sync",
        "Microsoft SCCM — device and patch management",
        "Oracle WebLogic — application runbook operations",
        "SMTP — email relay",
        "On-premise language model (GB10), with optional Azure OpenAI for frontier models",
        "Speech services — Whisper (speech-to-text) and Kokoro / Piper (text-to-speech)",
    ],
}

SOON = {
    "Service management (ITIL)": [
        "The whole solution running on real database records",
        "Incidents working end to end on the live system",
        "Automatic grouping of recurring incidents into a single problem",
        "Automatic change risk scoring from history, traffic and dependencies",
        "Maintenance-window conflict checking and safe-window suggestion",
        "Major-incident war room backed by live data and a real communications log",
    ],
    "Assistant and knowledge": [
        "Auto-drafted change-advisory (CAB) summary",
        "Executive summary regenerated live from the incident state",
    ],
    "Platform and operations": [
        "One-click seeding and reset of demonstration data",
        "Dashboards showing live metrics (mean time to resolve, SLA, ticket volume)",
    ],
    "System integrations": [],
}

# ── helpers ─────────────────────────────────────────────────────────────

def bottom_border(paragraph, color="2563EB", size="18"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    pPr.append(borders)


def run(p, text, *, size=10.5, bold=False, color=BODY, italic=False):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return r


def item_line(doc, mark, mark_color, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.space_after = Pt(3)
    run(p, f"{mark}  ", size=10.5, bold=True, color=mark_color)
    run(p, text, size=10.5, color=BODY)


# ── build ───────────────────────────────────────────────────────────────

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)

# Logo
if os.path.exists(LOGO):
    doc.add_picture(LOGO, width=Inches(1.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT

# Title block
t = doc.add_paragraph()
t.paragraph_format.space_before = Pt(10)
run(t, "AI Ops Platform", size=26, bold=True, color=NAVY)
s = doc.add_paragraph()
run(s, "Capabilities", size=15, bold=True, color=BLUE)
m = doc.add_paragraph()
run(m, "What the platform supports today, and what can be added within about 2 weeks.", size=11, color=GREY)
bottom_border(m)
doc.add_paragraph()

# Legend
lg = doc.add_paragraph()
run(lg, "✓  ", size=11, bold=True, color=GREEN)
run(lg, "Available today        ", size=11, color=BODY)
run(lg, "+  ", size=11, bold=True, color=BLUE)
run(lg, "Available within about 2 weeks", size=11, color=BODY)
doc.add_paragraph()

# Per-domain sections
for domain in DOMAINS:
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(8)
    run(h, domain, size=14, bold=True, color=NAVY)
    bottom_border(h, color="D1D5DB", size="8")
    for text in AVAILABLE.get(domain, []):
        item_line(doc, "✓", GREEN, text)
    for text in SOON.get(domain, []):
        item_line(doc, "+", BLUE, text)

# Footer note
doc.add_paragraph()
f = doc.add_paragraph()
run(f, "YourCompass — confidential.", size=9, italic=True, color=GREY)

doc.save(OUT)
print(f"wrote {OUT}")
